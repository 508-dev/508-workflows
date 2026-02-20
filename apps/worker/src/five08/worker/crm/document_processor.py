"""Resume document text extraction."""

import hashlib
import io
import logging
import re
from pathlib import Path

from five08.worker.config import settings

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Extract text from supported resume file formats."""

    def __init__(self) -> None:
        self.allowed_extensions = settings.allowed_file_extensions
        self.max_file_size = settings.max_file_size_mb * 1024 * 1024
        self._content_cache: dict[str, str] = {}

    def get_content_hash(self, content: bytes) -> str:
        """Hash bytes for extraction caching."""
        return hashlib.sha256(content).hexdigest()

    def is_valid_file(self, filename: str, file_size: int) -> tuple[bool, str | None]:
        """Validate extension and size."""
        if file_size > self.max_file_size:
            return False, f"File size {file_size} exceeds maximum {self.max_file_size}"

        ext = Path(filename).suffix.lower().lstrip(".")
        if ext not in self.allowed_extensions:
            return (
                False,
                f"File extension '{ext}' not allowed. Allowed: {self.allowed_extensions}",
            )
        return True, None

    def extract_text_from_docx(self, content: bytes) -> str:
        """Extract text from .docx content."""
        try:
            from docx import Document
        except Exception as exc:  # pragma: no cover - import failure is env-dependent
            raise ValueError(f"DOCX processing dependency missing: {exc}") from exc

        try:
            document = Document(io.BytesIO(content))
            chunks: list[str] = []

            for paragraph in document.paragraphs:
                text = paragraph.text.strip()
                if text:
                    chunks.append(text)

            for table in document.tables:
                for row in table.rows:
                    row_cells = [
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    ]
                    if row_cells:
                        chunks.append(" | ".join(row_cells))

            return "\n".join(chunks)
        except Exception as exc:
            logger.error("Error extracting DOCX text: %s", exc)
            raise ValueError(f"Failed to extract text from DOCX: {exc}") from exc

    def extract_text_from_pdf(self, content: bytes) -> str:
        """Extract text from PDF content."""
        try:
            from pdfminer.high_level import extract_text as extract_pdf_text
        except Exception as exc:  # pragma: no cover - import failure is env-dependent
            raise ValueError(f"PDF processing dependency missing: {exc}") from exc

        try:
            text = extract_pdf_text(io.BytesIO(content))
            return text.strip()
        except Exception as exc:
            logger.error("Error extracting PDF text: %s", exc)
            raise ValueError(f"Failed to extract text from PDF: {exc}") from exc

    def extract_text_from_doc(self, content: bytes) -> str:
        """Best-effort extraction from legacy .doc binary."""
        try:
            text = content.decode("utf-8", errors="ignore")
            text = re.sub(r"[^\x20-\x7E\n\r\t]", " ", text)
            text = re.sub(r"\s+", " ", text)
            return text.strip()
        except Exception as exc:
            logger.error("Error extracting DOC text: %s", exc)
            raise ValueError(f"Failed to extract text from DOC: {exc}") from exc

    def extract_text_from_txt(self, content: bytes) -> str:
        """Extract text from UTF-8 text files."""
        try:
            return content.decode("utf-8", errors="ignore").strip()
        except Exception as exc:
            logger.error("Error extracting TXT text: %s", exc)
            raise ValueError(f"Failed to extract text from TXT: {exc}") from exc

    def extract_text(self, content: bytes, filename: str) -> str:
        """Extract text from supported format and cache results."""
        content_hash = self.get_content_hash(content)
        if content_hash in self._content_cache:
            return self._content_cache[content_hash]

        is_valid, error = self.is_valid_file(filename, len(content))
        if not is_valid:
            raise ValueError(error or "Invalid file")

        ext = Path(filename).suffix.lower()
        if ext == ".pdf":
            text = self.extract_text_from_pdf(content)
        elif ext == ".docx":
            text = self.extract_text_from_docx(content)
        elif ext == ".doc":
            text = self.extract_text_from_doc(content)
        elif ext == ".txt":
            text = self.extract_text_from_txt(content)
        else:
            raise ValueError(f"Unsupported file type: {ext}")

        if not text.strip():
            raise ValueError("No text could be extracted from document")

        self._content_cache[content_hash] = text
        return text
