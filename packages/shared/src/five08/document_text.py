"""Shared document text extraction helpers."""

from __future__ import annotations

import io
import re

from five08.pdf import extract_pdf_text_with_links


def document_file_extension(filename: str | None) -> str:
    if not filename or "." not in filename:
        return ""
    return "." + filename.rsplit(".", 1)[-1].lower()


def extract_document_text(content: bytes, *, filename: str | None) -> str:
    extension = document_file_extension(filename)

    if extension == ".pdf":
        return extract_pdf_text_with_links(content).strip()
    if extension == ".docx":
        return _extract_docx_text(content).strip()
    if extension == ".doc":
        return _extract_doc_text(content).strip()
    return content.decode("utf-8", errors="ignore").strip()


def _extract_docx_text(content: bytes) -> str:
    try:
        from docx import Document
    except ModuleNotFoundError as exc:
        raise ValueError(
            "python-docx is required to extract text from .docx files; "
            "please install the 'python-docx' package."
        ) from exc

    try:
        document = Document(io.BytesIO(content))
        chunks: list[str] = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                chunks.append(text)

        for table in document.tables:
            for row in table.rows:
                row_cells = [
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                ]
                if row_cells:
                    chunks.append(" | ".join(row_cells))

        return "\n".join(chunks)
    except Exception as exc:
        raise ValueError("Failed to extract text from DOCX document.") from exc


def _extract_doc_text(content: bytes) -> str:
    text = content.decode("utf-8", errors="ignore")
    text = re.sub(r"[^\x20-\x7E\n\r\t]", " ", text)
    text = re.sub(r"\s+", " ", text)
    return text
